import os
import sys
import logging
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import img2pdf
from PIL import Image, ImageTk
import threading
import fitz
import openpyxl
import xlrd
from openpyxl import Workbook
import traceback
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.getLogger("pikepdf").setLevel(logging.ERROR)
Image.MAX_IMAGE_PIXELS = None
os.environ["PYDEVD_DISABLE_FILE_VALIDATION"] = "1"


def get_real_desktop_path():
    try:
        from win32com.shell import shell, shellcon
        desktop = shell.SHGetFolderPath(0, shellcon.CSIDL_DESKTOP, None, 0)
        if desktop and os.path.exists(desktop) and os.access(desktop, os.W_OK):
            return desktop
    except:
        pass

    user_profile = os.environ.get("USERPROFILE", "")
    if user_profile:
        desktop = os.path.join(user_profile, "Desktop")
        if os.path.exists(desktop) and os.access(desktop, os.W_OK):
            return desktop

    # 如果桌面不可用，使用文档文件夹
    documents = os.path.join(os.path.expanduser("~"), "Documents")
    if os.path.exists(documents) and os.access(documents, os.W_OK):
        return documents

    # 最后使用当前目录
    return os.path.dirname(sys.executable) if hasattr(sys, 'frozen') else os.getcwd()


def check_pdf_type(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        total_text = ""
        for page_num in range(min(3, len(doc))):
            page = doc[page_num]
            text = page.get_text().strip()
            total_text += text
            if len(total_text) > 100:
                doc.close()
                return "文本型"
        doc.close()
        return "扫描型"
    except Exception as e:
        print(f"PDF类型检测异常: {str(e)}")
        return "扫描型"


def clean_temp_files(file_list):
    for f in file_list:
        if os.path.exists(f):
            try:
                os.remove(f)
            except:
                pass


def text_pdf_to_word(input_file, output_file, progress_callback=None):
    try:
        pdf_doc = fitz.open(input_file)
        total_pages = len(pdf_doc)
        doc = Document()
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)

        if progress_callback:
            progress_callback(20)

        for page_num in range(total_pages):
            page = pdf_doc[page_num]
            text = page.get_text()

            if text.strip():
                paragraph = doc.add_paragraph(text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if page_num < total_pages - 1:
                    doc.add_page_break()

            if progress_callback:
                progress = 20 + (page_num + 1) / total_pages * 70
                progress_callback(progress)

        pdf_doc.close()
        doc.save(output_file)

        if progress_callback:
            progress_callback(100)
        return True, f"文本型PDF转换完成！共提取{total_pages}页文本"
    except Exception as e:
        error_msg = f"文本型PDF转换失败：{str(e)}"
        return False, error_msg


def scan_pdf_to_word(input_file, output_file, progress_callback=None):
    temp_files = []
    try:
        pdf_doc = fitz.open(input_file)
        total_pages = len(pdf_doc)
        img_paths = []
        temp_dir = os.path.join(os.environ.get("TEMP", ""), "pdf_converter_temp")
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        for idx in range(total_pages):
            page = pdf_doc[idx]
            zoom = 2.0
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix)
            img_path = os.path.join(temp_dir, f"temp_page_{idx + 1}.png")
            pix.save(img_path)
            img_paths.append(img_path)
            temp_files.append(img_path)

            if progress_callback:
                progress = (idx + 1) / total_pages * 50
                progress_callback(progress)

        pdf_doc.close()

        if progress_callback:
            progress_callback(55)

        doc = Document()
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        image_width = section.page_width - section.left_margin - section.right_margin

        for i, img_path in enumerate(img_paths):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(img_path, width=image_width)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if i < len(img_paths) - 1:
                doc.add_page_break()

        if progress_callback:
            progress_callback(85)

        doc.save(output_file)
        clean_temp_files(temp_files)

        if progress_callback:
            progress_callback(100)
        return True, f"扫描型PDF转换完成！共{total_pages}页已转为图片插入Word"
    except Exception as e:
        clean_temp_files(temp_files)
        error_msg = f"扫描型PDF转换失败：{str(e)}"
        return False, error_msg


def img_to_pdf(input_files, output_file, progress_callback=None):
    try:
        # 确保输入文件是字符串列表
        cleaned_files = []
        for f in input_files:
            if isinstance(f, tuple):
                cleaned_files.append(str(f[0]))
            elif isinstance(f, str):
                cleaned_files.append(f)
            else:
                cleaned_files.append(str(f))

        input_files = cleaned_files

        # 确保输出目录存在
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        img_list = []
        total_imgs = len(input_files)

        # 检查所有输入文件是否存在
        for img_path in input_files:
            if not os.path.exists(img_path):
                return False, f"图片文件不存在：{img_path}"

        for idx, img_path in enumerate(input_files):
            try:
                with Image.open(img_path) as img:
                    # 转换RGBA为RGB
                    if img.mode in ("RGBA", "P"):
                        img = img.convert("RGB")
                    # 复制图像数据，确保它在内存中保持活跃
                    img_copy = img.copy()
                    img_list.append(img_copy)
            except Exception as e:
                return False, f"无法打开图片文件 {img_path}：{str(e)}"

            if progress_callback:
                progress = (idx + 1) / total_imgs * 90
                progress_callback(progress)

        # 保存为PDF
        if img_list:
            try:
                # 使用临时文件方式避免直接写入失败
                temp_output = output_file + ".tmp"
                if len(img_list) == 1:
                    img_list[0].save(temp_output, "PDF", resolution=100.0)
                else:
                    img_list[0].save(
                        temp_output,
                        "PDF",
                        save_all=True,
                        append_images=img_list[1:],
                        resolution=100.0
                    )
                
                # 如果临时文件创建成功，则替换原文件
                if os.path.exists(temp_output):
                    if os.path.exists(output_file):
                        os.remove(output_file)
                    os.rename(temp_output, output_file)
            except Exception as e:
                # 尝试使用img2pdf作为备用方案
                try:
                    import img2pdf
                    with open(output_file, "wb") as f:
                        f.write(img2pdf.convert(input_files))
                except Exception as e2:
                    # 如果备用方案也失败，返回原始错误
                    if os.path.exists(output_file + ".tmp"):
                        try:
                            os.remove(output_file + ".tmp")
                        except:
                            pass
                    return False, f"保存PDF文件失败：{str(e)}，备用方案也失败：{str(e2)}"
        else:
            return False, "没有有效的图片可转换"

        if progress_callback:
            progress_callback(100)
        return True, f"图片转PDF完成！共转换{len(img_list)}张图片"
    except Exception as e:
        error_msg = f"图片转PDF失败：{str(e)}"
        traceback_info = traceback.format_exc()
        return False, f"{error_msg}\n\n详细错误：{traceback_info}"


def word_to_pdf(input_file, output_file, progress_callback=None):
    try:
        temp_dir = os.path.join(os.environ.get("TEMP", ""), "word_to_pdf_temp")
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
        try:
            doc = Document(input_file)
            pdf_doc = fitz.open()

            for para in doc.paragraphs:
                if para.text.strip():
                    page = pdf_doc.new_page(width=595, height=842)
                    text = para.text
                    page.insert_text((50, 50), text, fontsize=12)

            if len(pdf_doc) == 0:
                pdf_doc.new_page(width=595, height=842)

            pdf_doc.save(output_file)
            pdf_doc.close()

            if progress_callback:
                progress_callback(100)
            return True, "Word转PDF完成（文本内容转换）"
        except Exception as e:
            return word_to_img_then_pdf(input_file, output_file, progress_callback)
    except Exception as e:
        error_msg = f"Word转PDF失败：{str(e)}"
        return False, error_msg


def word_to_img_then_pdf(input_file, output_file, progress_callback=None):
    try:
        temp_dir = os.path.join(os.environ.get("TEMP", ""), "word_temp")
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        success, msg = word_to_img(input_file, temp_dir, progress_callback)

        if not success:
            return False, f"Word转PDF失败：{msg}"

        img_files = []
        for f in os.listdir(temp_dir):
            if f.endswith(('.png', '.jpg', '.jpeg')):
                img_files.append(os.path.join(temp_dir, f))

        if not img_files:
            return False, "Word转PDF失败：未生成图片文件"

        img_list = []
        for img_path in img_files:
            with Image.open(img_path) as img:
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img_list.append(img)

        if img_list:
            img_list[0].save(
                output_file,
                "PDF",
                save_all=True,
                append_images=img_list[1:] if len(img_list) > 1 else [],
                quality=95
            )

        for f in img_files:
            if os.path.exists(f):
                os.remove(f)
        if os.path.exists(temp_dir):
            os.rmdir(temp_dir)

        return True, f"Word转PDF完成！共转换{len(img_files)}页"
    except Exception as e:
        return False, f"Word转PDF失败：{str(e)}"


def word_to_img(input_file, output_dir, progress_callback=None):
    try:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        doc = Document(input_file)
        paragraphs = list(doc.paragraphs)
        total_paras = len(paragraphs)

        if total_paras == 0:
            img = Image.new('RGB', (800, 600), color='white')
            img_path = os.path.join(output_dir, "word_page_1.png")
            img.save(img_path)

            if progress_callback:
                progress_callback(100)
            return True, f"Word转图片完成！保存至：{output_dir}"

        from PIL import ImageDraw, ImageFont

        for idx, para in enumerate(paragraphs):
            if para.text.strip():
                img = Image.new('RGB', (800, 200), color='white')
                draw = ImageDraw.Draw(img)
                try:
                    font = ImageFont.truetype("arial.ttf", 14)
                except:
                    font = ImageFont.load_default()
                draw.text((10, 10), para.text, fill='black', font=font)

                img_path = os.path.join(output_dir, f"word_page_{idx + 1}.png")
                img.save(img_path)

            if progress_callback:
                progress = (idx + 1) / total_paras * 100
                progress_callback(progress)

        return True, f"Word转图片完成！共生成{total_paras}张图片，保存至：{output_dir}"
    except Exception as e:
        error_msg = f"Word转图片失败：{str(e)}"
        return False, error_msg


def pdf_to_img(input_file, output_dir, progress_callback=None):
    try:
        pdf_doc = fitz.open(input_file)
        total_pages = len(pdf_doc)

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        for page_num in range(total_pages):
            page = pdf_doc[page_num]
            pix = page.get_pixmap(dpi=150)
            img_path = os.path.join(output_dir, f"pdf_page_{page_num + 1}.png")
            pix.save(img_path)

            if progress_callback:
                progress = (page_num + 1) / total_pages * 100
                progress_callback(progress)

        pdf_doc.close()
        return True, f"PDF转图片完成！共生成{total_pages}张图片，保存至：{output_dir}"
    except Exception as e:
        return False, f"PDF转图片失败：{str(e)}"


def xls_to_word(input_file, output_file, progress_callback=None):
    try:
        data = []
        sheet_names = []

        if input_file.endswith(".xls"):
            wb = xlrd.open_workbook(input_file)
            sheet_names = wb.sheet_names()
            for sheet_name in sheet_names:
                ws = wb.sheet_by_name(sheet_name)
                sheet_data = []
                for row_idx in range(ws.nrows):
                    row_data = [ws.cell_value(row_idx, col_idx) for col_idx in range(ws.ncols)]
                    sheet_data.append(row_data)
                data.append((sheet_name, sheet_data))
            wb.release_resources()
        else:
            wb = openpyxl.load_workbook(input_file, data_only=True)
            sheet_names = wb.sheetnames
            for sheet_name in sheet_names:
                ws = wb[sheet_name]
                sheet_data = []
                for row in ws.iter_rows(values_only=True):
                    sheet_data.append([cell if cell is not None else "" for cell in row])
                data.append((sheet_name, sheet_data))
            wb.close()

        if progress_callback:
            progress_callback(40)

        doc = Document()
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        doc.add_heading('Excel数据转换结果', 0)

        total_sheets = len(data)
        for sheet_idx, (sheet_name, sheet_data) in enumerate(data):
            doc.add_heading(f'工作表：{sheet_name}', level=1)

            if sheet_data:
                table = doc.add_table(rows=len(sheet_data), cols=len(sheet_data[0]))
                table.style = 'Table Grid'

                for row_idx, row_data in enumerate(sheet_data):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell_value in enumerate(row_data):
                        row_cells[col_idx].text = str(cell_value)

            if sheet_idx < total_sheets - 1:
                doc.add_page_break()

            if progress_callback:
                progress = 40 + (sheet_idx + 1) / total_sheets * 50
                progress_callback(progress)

        doc.save(output_file)

        if progress_callback:
            progress_callback(100)
        return True, f"XLS转Word完成！共处理{total_sheets}个工作表，{len(data[0][1]) if data else 0}行数据"
    except Exception as e:
        error_msg = f"XLS转Word失败：{str(e)}"
        return False, error_msg


def show_image_grid_dialog(files, file_path, root):
    """显示图片网格预览对话框，支持拖动排序和删除"""
    grid_dialog = tk.Toplevel(root)
    grid_dialog.title("图片网格预览与排序")
    grid_dialog.geometry("1000x700")
    grid_dialog.configure(bg="#f8f9fa")
    
    # 居中显示
    grid_dialog.update_idletasks()
    w = grid_dialog.winfo_width()
    h = grid_dialog.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (w // 2)
    y = (root.winfo_screenheight() // 2) - (h // 2)
    grid_dialog.geometry(f"{w}x{h}+{x}+{y}")

    # 存储文件路径和对应的预览图片
    file_data = []
    for f in files:
        file_data.append({
            'path': str(f),
            'thumbnail': None,
            'label': None  # 用于存储tkinter Label组件
        })

    # 创建主框架
    main_frame = tk.Frame(grid_dialog, bg="#f8f9fa")
    main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

    # 创建画布和滚动条
    canvas_frame = tk.Frame(main_frame)
    canvas_frame.pack(fill=tk.BOTH, expand=True)

    canvas = tk.Canvas(canvas_frame, bg="white")
    v_scrollbar = tk.Scrollbar(canvas_frame, orient=tk.VERTICAL, command=canvas.yview)
    h_scrollbar = tk.Scrollbar(main_frame, orient=tk.HORIZONTAL, command=canvas.xview)
    
    canvas.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)

    v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

    # 创建预览框架
    preview_frame = tk.Frame(canvas, bg="white")
    canvas.create_window((0, 0), window=preview_frame, anchor=tk.NW)

    # 用于存储拖动信息
    drag_data = {"widget": None, "start_x": 0, "start_y": 0, "index": None}

    def load_thumbnails():
        """加载所有图片的缩略图"""
        for i, data in enumerate(file_data):
            try:
                img = Image.open(data['path'])
                # 创建缩略图
                img.thumbnail((150, 150), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(img)
                data['thumbnail'] = photo
            except Exception as e:
                print(f"无法加载缩略图: {data['path']}, 错误: {e}")

    def create_preview_grid():
        #创建预览网格
        # 清除现有组件
        for widget in preview_frame.winfo_children():
            widget.destroy()
        
        # 每行显示6张图片
        cols = 6
        for i, data in enumerate(file_data):
            row = i // cols
            col = i % cols
            
            # 创建包含图片和控制按钮的框架
            img_frame = tk.Frame(preview_frame, relief=tk.RAISED, bd=2, bg="white")
            img_frame.grid(row=row, column=col, padx=10, pady=10)
            
            # 显示缩略图
            if data['thumbnail']:
                img_label = tk.Label(img_frame, image=data['thumbnail'], bg="white", relief=tk.SUNKEN)
            else:
                img_label = tk.Label(img_frame, text="图片\n加载失败", width=20, height=10, 
                                   bg="lightgray", relief=tk.SUNKEN)
            
            img_label.grid(row=0, column=0, columnspan=2)
            data['label'] = img_label  # 保存引用用于拖动
            
            # 文件名
            filename = os.path.basename(data['path'])
            name_label = tk.Label(img_frame, text=filename[:15] + "..." if len(filename) > 15 else filename, 
                                font=("微软雅黑", 8), bg="white")
            name_label.grid(row=1, column=0, columnspan=2, pady=(5, 0))
            
            # 删除按钮
            delete_btn = tk.Button(img_frame, text="🗑", font=("微软雅黑", 8), 
                                 command=lambda idx=i: delete_image(idx),
                                 bg="#dc3545", fg="white", width=3)
            delete_btn.grid(row=2, column=0, padx=2, pady=5, sticky=tk.E)
            
            # 上移按钮
            up_btn = tk.Button(img_frame, text="⬆", font=("微软雅黑", 8), 
                             command=lambda idx=i: move_image_up(idx),
                             bg="#ffc107", fg="black", width=3)
            up_btn.grid(row=2, column=1, padx=2, pady=5, sticky=tk.W)
            
            # 绑定拖动事件
            img_label.bind("<Button-1>", lambda e, idx=i: start_drag(e, idx))
            img_label.bind("<B1-Motion>", drag_motion)
            img_label.bind("<ButtonRelease-1>", lambda e, idx=i: stop_drag(e, idx))

    def start_drag(event, index):
        """开始拖动"""
        drag_data["widget"] = file_data[index]['label']
        drag_data["start_x"] = event.x
        drag_data["start_y"] = event.y
        drag_data["index"] = index
        # 提升到顶层
        drag_data["widget"].lift()

    def drag_motion(event):
        """拖动中"""
        if drag_data["widget"]:
            x = drag_data["widget"].winfo_x() - drag_data["start_x"] + event.x
            y = drag_data["widget"].winfo_y() - drag_data["start_y"] + event.y
            # 由于Label在Frame内，我们需要使用place来精确定位
            drag_data["widget"].master.update()  # 更新布局
            drag_data["widget"].place(x=x, y=y)
            drag_data["widget"].lift()

    def stop_drag(event, index):
        """停止拖动"""
        if drag_data["widget"]:
            # 计算当前鼠标位置在网格中的目标位置
            canvas_x = canvas.canvasx(event.x)
            canvas_y = canvas.canvasy(event.y)
            
            # 获取所有图片框架的位置
            target_idx = find_drop_target(canvas_x, canvas_y, index)
            if target_idx is not None and target_idx != index:
                # 交换图片位置
                file_data[index], file_data[target_idx] = file_data[target_idx], file_data[index]
                create_preview_grid()  # 重新创建网格
                
            drag_data["widget"] = None
            drag_data["index"] = None
            # 重新加载网格以反映新顺序

    def find_drop_target(x, y, source_idx):
        """找到拖放目标位置"""
        min_dist = float('inf')
        target_idx = None
        
        cols = 6
        for i, data in enumerate(file_data):
            if i == source_idx:
                continue
            row = i // cols
            col = i % cols
            # 计算每个框架的大致位置
            frame_x = col * 170 + 85  # 170是框架宽度+间距，85是中心点
            frame_y = row * 220 + 110  # 220是框架高度+间距，110是中心点
            
            dist = ((x - frame_x) ** 2 + (y - frame_y) ** 2) ** 0.5
            if dist < min_dist:
                min_dist = dist
                target_idx = i
        
        # 如果距离足够近，则认为是有效拖放
        if min_dist < 100:
            return target_idx
        return None

    def delete_image(index):
        """删除指定索引的图片"""
        if messagebox.askyesno("确认删除", f"确定要删除图片 {os.path.basename(file_data[index]['path'])} 吗？"):
            file_data.pop(index)
            create_preview_grid()

    def move_image_up(index):
        """上移图片"""
        if index > 0:
            file_data[index], file_data[index-1] = file_data[index-1], file_data[index]
            create_preview_grid()

    def confirm_order():
        """确认排序"""
        ordered_paths = [data['path'] for data in file_data]
        root.file_list = ordered_paths
        
        # 更新主界面显示
        if len(ordered_paths) == 1:
            display = f"✅ 已选中：\n{ordered_paths[0]}"
        else:
            display = f"✅ 已选中 {len(ordered_paths)} 个图片文件（已排序）\n"
            for i, f in enumerate(ordered_paths[:3]):
                display += f"{i + 1}. {os.path.basename(f)}\n"
            if len(ordered_paths) > 3:
                display += f"...等{len(ordered_paths)}个文件"
        
        file_path.set(display)
        grid_dialog.destroy()

    def cancel_order():
        """取消排序"""
        grid_dialog.destroy()

    # 加载缩略图并创建网格
    load_thumbnails()
    create_preview_grid()

    # 更新滚动区域
    preview_frame.update_idletasks()
    # 为Canvas设置滚动区域
    canvas.config(scrollregion=canvas.bbox("all"))

    # 按钮框架
    button_frame = tk.Frame(grid_dialog, bg="#f8f9fa")
    button_frame.pack(pady=10)

    btn_confirm = tk.Button(button_frame, text="✅ 确认排序", command=confirm_order,
                            font=("微软雅黑", 10, "bold"), bg="#0d6efd", fg="white",
                            padx=20, pady=5)
    btn_confirm.pack(side=tk.LEFT, padx=10)

    btn_cancel = tk.Button(button_frame, text="❌ 取消", command=cancel_order,
                           font=("微软雅黑", 10), bg="#dc3545", fg="white",
                           padx=15, pady=5)
    btn_cancel.pack(side=tk.LEFT, padx=10)

    # 刷新按钮
    btn_refresh = tk.Button(button_frame, text="🔄 刷新", command=create_preview_grid,
                           font=("微软雅黑", 10), bg="#28a745", fg="white",
                           padx=15, pady=5)
    btn_refresh.pack(side=tk.LEFT, padx=10)


def show_image_order_dialog(files, file_path, root):
    """显示图片排序对话框（保留原有的列表排序方式）"""
    # 创建新对话框使用网格预览
    show_image_grid_dialog(files, file_path, root)


def manage_file_list(converter_type, file_path, root):
    """管理文件列表，包括预览、删除和排序功能"""
    
    def preview_image(file_path):
        """预览图片文件"""
        try:
            img = Image.open(file_path)
            
            # 创建预览窗口
            preview_window = tk.Toplevel(root)
            preview_window.title(f"预览: {os.path.basename(file_path)}")
            preview_window.geometry("600x500")
            preview_window.configure(bg="#f8f9fa")
            
            # 调整图片大小以适应窗口
            img_width, img_height = img.size
            max_width, max_height = 550, 400
            
            scale = min(max_width/img_width, max_height/img_height)
            if scale < 1:
                new_size = (int(img_width*scale), int(img_height*scale))
                img = img.resize(new_size, Image.Resampling.LANCZOS)
            
            # 转换为tkinter可用的格式
            photo = ImageTk.PhotoImage(img)
            
            # 创建画布显示图片
            canvas = tk.Canvas(preview_window, width=min(img_width, max_width), 
                              height=min(img_height, max_height), bg="white")
            canvas.pack(pady=10, padx=10)
            canvas.create_image(0, 0, anchor=tk.NW, image=photo)
            
            # 保持图片引用防止被垃圾回收
            canvas.image = photo
            
            # 添加关闭按钮
            close_btn = tk.Button(preview_window, text="关闭", command=preview_window.destroy,
                                 font=("微软雅黑", 10), bg="#0d6efd", fg="white")
            close_btn.pack(pady=10)
            
        except Exception as e:
            messagebox.showerror("预览错误", f"无法预览图片: {str(e)}")
    
    def delete_selected():
        """删除选中的文件"""
        selected_indices = listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("删除", "请先选择要删除的文件")
            return
        
        # 从后往前删除，避免索引变化
        for index in reversed(selected_indices):
            root.file_list.pop(index)
        
        update_list_display()
    
    def move_up():
        """上移选中的文件"""
        selected_indices = listbox.curselection()
        if not selected_indices or selected_indices[0] == 0:
            return
        
        index = selected_indices[0]
        # 交换元素
        root.file_list[index], root.file_list[index-1] = root.file_list[index-1], root.file_list[index]
        
        update_list_display()
        listbox.selection_set(index-1)
    
    def move_down():
        """下移选中的文件"""
        selected_indices = listbox.curselection()
        if not selected_indices or selected_indices[0] >= len(root.file_list)-1:
            return
        
        index = selected_indices[0]
        # 交换元素
        root.file_list[index], root.file_list[index+1] = root.file_list[index+1], root.file_list[index]
        
        update_list_display()
        listbox.selection_set(index+1)
    
    def update_list_display():
        """更新列表显示"""
        listbox.delete(0, tk.END)
        for i, file_path in enumerate(root.file_list):
            filename = os.path.basename(file_path)
            listbox.insert(tk.END, f"{i+1}. {filename}")
        
        # 更新显示文本
        if len(root.file_list) == 0:
            file_path.set("")
        elif len(root.file_list) == 1:
            file_path.set(f"✅ 已选中：\n{os.path.basename(root.file_list[0])}")
        else:
            file_path.set(f"✅ 已选中 {len(root.file_list)} 个文件")
    
    def confirm_selection():
        """确认选择"""
        manager_window.destroy()
    
    # 创建管理窗口
    manager_window = tk.Toplevel(root)
    manager_window.title("文件管理")
    manager_window.geometry("700x500")
    manager_window.configure(bg="#f8f9fa")
    
    # 居中显示
    manager_window.update_idletasks()
    w = manager_window.winfo_width()
    h = manager_window.winfo_height()
    x = (manager_window.winfo_screenwidth() // 2) - (w // 2)
    y = (manager_window.winfo_screenheight() // 2) - (h // 2)
    manager_window.geometry(f"{w}x{h}+{x}+{y}")
    
    tk.Label(manager_window, text="📁 已选择的文件", font=("微软雅黑", 12, "bold"),
             bg="#f8f9fa", fg="#0d6efd").pack(pady=10)
    
    # 文件列表
    list_frame = tk.Frame(manager_window)
    list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
    
    scrollbar = tk.Scrollbar(list_frame)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
    listbox = tk.Listbox(list_frame, selectmode=tk.SINGLE, yscrollcommand=scrollbar.set,
                         font=("微软雅黑", 10), height=15, bg="white", fg="#333333")
    listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    scrollbar.config(command=listbox.yview)
    
    # 填充列表
    update_list_display()
    
    # 按钮框架
    button_frame = tk.Frame(manager_window, bg="#f8f9fa")
    button_frame.pack(pady=15)
    
    # 预览按钮
    preview_btn = tk.Button(button_frame, text="👁 预览", command=lambda: (
        preview_image(root.file_list[listbox.curselection()[0]]) if listbox.curselection() else 
        messagebox.showwarning("预览", "请先选择要预览的文件")
    ), font=("微软雅黑", 10), bg="#28a745", fg="white", padx=15, pady=5)
    preview_btn.pack(side=tk.LEFT, padx=5)
    
    # 删除按钮
    delete_btn = tk.Button(button_frame, text="🗑 删除", command=delete_selected,
                          font=("微软雅黑", 10), bg="#dc3545", fg="white", padx=15, pady=5)
    delete_btn.pack(side=tk.LEFT, padx=5)
    
    # 上移按钮
    up_btn = tk.Button(button_frame, text="↑ 上移", command=move_up,
                       font=("微软雅黑", 10), bg="#ffc107", fg="#212529", padx=15, pady=5)
    up_btn.pack(side=tk.LEFT, padx=5)
    
    # 下移按钮
    down_btn = tk.Button(button_frame, text="↓ 下移", command=move_down,
                         font=("微软雅黑", 10), bg="#ffc107", fg="#212529", padx=15, pady=5)
    down_btn.pack(side=tk.LEFT, padx=5)
    
    # 确认按钮
    confirm_btn = tk.Button(button_frame, text="✅ 确认", command=confirm_selection,
                           font=("微软雅黑", 10, "bold"), bg="#0d6efd", fg="white", padx=20, pady=5)
    confirm_btn.pack(side=tk.LEFT, padx=20)
    
    # 添加文件按钮
    def add_more_files():
        title_map = {
            "pdf2word": "选择要转换的PDF文件",
            "word2pdf": "选择要转换的Word文件",
            "img2pdf": "选择要转换的图片文件（可多选）",
            "word2img": "选择要转换的Word文件",
            "pdf2img": "选择要转换的PDF文件",
            "xls2word": "选择要转换的Excel文件"
        }
        type_map = {
            "pdf2word": [("PDF文件", "*.pdf"), ("所有文件", "*.*")],
            "word2pdf": [("Word文件", "*.docx;*.doc"), ("所有文件", "*.*")],
            "img2pdf": [("图片文件", "*.jpg;*.jpeg;*.png;*.bmp;*.gif"), ("所有文件", "*.*")],
            "word2img": [("Word文件", "*.docx;*.doc"), ("所有文件", "*.*")],
            "pdf2img": [("PDF文件", "*.pdf"), ("所有文件", "*.*")],
            "xls2word": [("Excel文件", "*.xls;*.xlsx"), ("所有文件", "*.*")]
        }
        
        files = filedialog.askopenfilenames(title=title_map[converter_type], filetypes=type_map[converter_type])
        
        if files:
            root.file_list.extend([str(f) for f in files])
            update_list_display()
    
    add_btn = tk.Button(button_frame, text="➕ 添加", command=add_more_files,
                        font=("微软雅黑", 10), bg="#17a2b8", fg="white", padx=15, pady=5)
    add_btn.pack(side=tk.LEFT, padx=5)


def select_file(converter_type, file_path, root):
    file_path.set("")
    title_map = {
        "pdf2word": "选择要转换的PDF文件",
        "word2pdf": "选择要转换的Word文件",
        "img2pdf": "选择要转换的图片文件（可多选）",
        "word2img": "选择要转换的Word文件",
        "pdf2img": "选择要转换的PDF文件",
        "xls2word": "选择要转换的Excel文件"
    }
    type_map = {
        "pdf2word": [("PDF文件", "*.pdf"), ("所有文件", "*.*")],
        "word2pdf": [("Word文件", "*.docx;*.doc"), ("所有文件", "*.*")],
        "img2pdf": [("图片文件", "*.jpg;*.jpeg;*.png;*.bmp;*.gif"), ("所有文件", "*.*")],
        "word2img": [("Word文件", "*.docx;*.doc"), ("所有文件", "*.*")],
        "pdf2img": [("PDF文件", "*.pdf"), ("所有文件", "*.*")],
        "xls2word": [("Excel文件", "*.xls;*.xlsx"), ("所有文件", "*.*")]
    }

    files = filedialog.askopenfilenames(title=title_map[converter_type], filetypes=type_map[converter_type])

    if files:
        # 对于图片转PDF且选择了多个文件，显示排序界面
        if converter_type == "img2pdf" and len(files) > 1:
            show_image_order_dialog(files, file_path, root)
        else:
            root.file_list = [str(f) for f in files]
            display = f"✅ 已选中：\n{files[0]}" if len(files) == 1 else f"✅ 已选中 {len(files)} 个文件"
            file_path.set(display)
            
            # 如果是图片文件，提供管理选项
            if converter_type in ["img2pdf"] and len(files) > 1:
                if messagebox.askyesno("文件管理", f"已选择 {len(files)} 个文件，是否需要管理（预览/排序/删除）？"):
                    manage_file_list(converter_type, file_path, root)


def update_progress(value, progress_var, root):
    progress_var.set(value)
    root.update_idletasks()


def convert_thread(converter_type, input_files, new_name, root, progress_var, convert_btn, name_entry, file_path):
    desktop = get_real_desktop_path()

    # 确保桌面路径存在且可写
    if not os.path.exists(desktop) or not os.access(desktop, os.W_OK):
        # 如果桌面不可用，尝试其他常见路径
        fallback_paths = [
            os.path.join(os.path.expanduser("~"), "Documents"),
            os.path.join(os.path.expanduser("~"), "Downloads"),
            os.path.dirname(sys.executable) if hasattr(sys, 'frozen') else os.getcwd(),
            os.getcwd()
        ]
        
        desktop = None
        for path in fallback_paths:
            if os.path.exists(path) and os.access(path, os.W_OK):
                desktop = path
                break
        
        # 如果仍没有可用路径，创建一个临时目录
        if desktop is None:
            import tempfile
            desktop = tempfile.mkdtemp(prefix="file_converter_")
            
        root.after(0, lambda: messagebox.showwarning("权限提示", f"桌面不可写，保存到：\n{desktop}"))

    success = False
    msg = ""
    output = ""

    try:
        if converter_type == "pdf2word":
            selected_pdf = input_files[0]
            pdf_type = check_pdf_type(selected_pdf)
            root.after(0, lambda: messagebox.showinfo("PDF类型检测", f"检测到PDF类型：{pdf_type}\n开始转换..."))
            output = os.path.join(desktop, f"{new_name}.docx")

            if pdf_type == "文本型":
                success, msg = text_pdf_to_word(
                    selected_pdf,
                    output,
                    lambda v: update_progress(v, progress_var, root)
                )
            else:
                success, msg = scan_pdf_to_word(
                    selected_pdf,
                    output,
                    lambda v: update_progress(v, progress_var, root)
                )

        elif converter_type == "word2pdf":
            selected_word = input_files[0]
            output = os.path.join(desktop, f"{new_name}.pdf")
            success, msg = word_to_pdf(
                selected_word,
                output,
                lambda v: update_progress(v, progress_var, root)
            )

        elif converter_type == "img2pdf":
            if not new_name.lower().endswith('.pdf'):
                new_name = f"{new_name}.pdf"

            output = os.path.join(desktop, new_name)

            success, msg = img_to_pdf(
                input_files,
                output,
                lambda v: update_progress(v, progress_var, root)
            )

        elif converter_type == "word2img":
            selected_word = input_files[0]
            output = os.path.join(desktop, new_name)
            success, msg = word_to_img(
                selected_word,
                output,
                lambda v: update_progress(v, progress_var, root)
            )

        elif converter_type == "pdf2img":
            selected_pdf = input_files[0]
            output = os.path.join(desktop, new_name)
            success, msg = pdf_to_img(
                selected_pdf,
                output,
                lambda v: update_progress(v, progress_var, root)
            )

        elif converter_type == "xls2word":
            selected_xls = input_files[0]
            output = os.path.join(desktop, f"{new_name}.docx")
            success, msg = xls_to_word(
                selected_xls,
                output,
                lambda v: update_progress(v, progress_var, root)
            )

        else:
            msg = "无效转换类型"

    except Exception as e:
        error_trace = traceback.format_exc()
        msg = f"执行异常：{type(e).__name__} - {str(e)}\n\n详细错误：{error_trace}"

    root.after(0, lambda: finish_convert(success, msg, output, root, convert_btn, progress_var, name_entry, file_path))


def finish_convert(success, msg, output, root, convert_btn, progress_var, name_entry, file_path):
    convert_btn.config(state=tk.NORMAL, text="🚀 开始转换")
    progress_var.set(0)

    if success:
        messagebox.showinfo("转换成功 🎉", f"{msg}")
        name_entry.delete(0, tk.END)
        file_path.set("")
        root.file_list = []
    else:
        error_msg = f"转换失败 ❌\n\n{msg}"
        messagebox.showerror("错误", error_msg)


def start_convert(var, root, progress_var, convert_btn, name_entry, file_path):
    converter_type = var.get()
    input_files = getattr(root, "file_list", [])
    new_name = name_entry.get().strip()

    if not input_files:
        messagebox.showwarning("提示", "请先选择文件！")
        return
    if not new_name:
        messagebox.showwarning("提示", "请输入文件名！")
        return
    if any(char in new_name for char in r'\/:*?"<>|'):
        messagebox.showwarning("提示", "文件名不能包含：\\ / : * ? \" < > |")
        return

    # 检查文件是否存在
    for file in input_files:
        file_str = str(file)
        if not os.path.exists(file_str):
            messagebox.showerror("错误", f"选中的文件不存在：\n{file_str}")
            return

    convert_btn.config(state=tk.DISABLED, text="⏳ 转换中...")
    progress_var.set(0)

    t = threading.Thread(
        target=convert_thread,
        args=(converter_type, input_files, new_name, root, progress_var, convert_btn, name_entry, file_path),
        daemon=True
    )
    t.start()


if __name__ == "__main__":
    root = tk.Tk()
    root.title("📄 FMgaic")
    root.geometry("900x580")
    root.resizable(False, False)
    root.configure(bg="#f8f9fa")


    def center_win():
        root.update_idletasks()
        w = root.winfo_width()
        h = root.winfo_height()
        x = (root.winfo_screenwidth() // 2) - (w // 2)
        y = (root.winfo_screenheight() // 2) - (h // 2)
        root.geometry(f"{w}x{h}+{x}+{y}")


    center_win()

    progress_var = tk.DoubleVar(value=0)
    var = tk.StringVar(value="pdf2word")
    file_path = tk.StringVar(value="")
    root.file_list = []

    style = ttk.Style(root)
    style.theme_use("clam")
    style.configure("Modern.TButton", font=("微软雅黑", 11, "bold"),
                    background="#0d6efd", foreground="white", relief=tk.FLAT, padding=8)
    style.map("Modern.TButton",
              background=[("active", "#0b5ed7"), ("disabled", "#cccccc")],
              foreground=[("active", "white"), ("disabled", "#666666")])
    style.configure("Big.TButton", font=("微软雅黑", 14, "bold"),
                    background="#0d6efd", foreground="white", relief=tk.FLAT, padding=(30, 12))
    style.configure("Modern.TEntry", font=("微软雅黑", 11), padding=8,
                    relief=tk.FLAT, fieldbackground="white", foreground="#333333")
    style.configure("Modern.TRadiobutton", font=("微软雅黑", 11),
                    foreground="#0d6efd", background="#f8f9fa", padding=6)

    title_frame = tk.Frame(root, bg="#f8f9fa")
    title_frame.pack(fill=tk.X, pady=(10, 5), padx=40)
    title_label = tk.Label(title_frame, text="FMgaic",
                           font=("微软雅黑", 22, "bold"), fg="#0d6efd", bg="#f8f9fa")
    title_label.pack(expand=True)

    type_card = tk.Frame(root, bg="white", bd=0, relief=tk.FLAT)
    type_card.pack(fill=tk.X, padx=40, pady=8)
    type_title = tk.Label(type_card, text="📌 选择转换类型",
                          font=("微软雅黑", 12, "bold"), fg="#0d6efd", bg="white",
                          anchor="w", padx=20, pady=8)
    type_title.pack(fill=tk.X)

    r_frame = tk.Frame(type_card, bg="white", padx=20, pady=5)
    r_frame.pack(fill=tk.X)

    row1 = tk.Frame(r_frame, bg="white")
    row1.pack(fill=tk.X, pady=2)

    for i, (text, value) in enumerate(
            [("PDF → Word", "pdf2word"), ("Word → PDF", "word2pdf"), ("图片 → PDF", "img2pdf")]):
        cell_frame = tk.Frame(row1, bg="white")
        cell_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2)

        radio = ttk.Radiobutton(cell_frame, text=text, variable=var, value=value,
                                style="Modern.TRadiobutton")
        radio.pack(fill=tk.X, padx=5, pady=3)

    row2 = tk.Frame(r_frame, bg="white")
    row2.pack(fill=tk.X, pady=2)

    for i, (text, value) in enumerate(
            [("Word → 图片", "word2img"), ("PDF → 图片", "pdf2img"), ("XLS → Word", "xls2word")]):
        cell_frame = tk.Frame(row2, bg="white")
        cell_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=2)

        radio = ttk.Radiobutton(cell_frame, text=text, variable=var, value=value,
                                style="Modern.TRadiobutton")
        radio.pack(fill=tk.X, padx=5, pady=3)

    file_card = tk.Frame(root, bg="white", bd=0, relief=tk.FLAT)
    file_card.pack(fill=tk.X, padx=40, pady=8)
    file_title = tk.Label(file_card, text="📂 选择文件",
                          font=("微软雅黑", 12, "bold"), fg="#0d6efd", bg="white",
                          anchor="w", padx=20, pady=8)
    file_title.pack(fill=tk.X)
    file_inner = tk.Frame(file_card, bg="white", padx=20, pady=5)
    file_inner.pack(fill=tk.X)
    file_entry = ttk.Entry(file_inner, textvariable=file_path, state="readonly",
                           style="Modern.TEntry", width=80)
    file_entry.pack(side=tk.LEFT, padx=(0, 15), fill=tk.X, expand=True)
    select_btn = ttk.Button(file_inner, text="Go",
                            command=lambda: select_file(var.get(), file_path, root),
                            style="Modern.TButton", width=10)
    select_btn.pack(side=tk.RIGHT)

    name_card = tk.Frame(root, bg="white", bd=0, relief=tk.FLAT)
    name_card.pack(fill=tk.X, padx=40, pady=8)
    name_title = tk.Label(name_card, text="✏️ 自定义文件名/目录名",
                          font=("微软雅黑", 12, "bold"), fg="#0d6efd", bg="white",
                          anchor="w", padx=20, pady=8)
    name_title.pack(fill=tk.X)
    name_inner = tk.Frame(name_card, bg="white", padx=20, pady=5)
    name_inner.pack(fill=tk.X)
    name_entry = ttk.Entry(name_inner, style="Modern.TEntry", width=80)
    name_entry.pack(side=tk.LEFT, padx=(0, 15), fill=tk.X, expand=True)

    progress_frame = tk.Frame(root, bg="#f8f9fa")
    progress_frame.pack(fill=tk.X, padx=40, pady=15)
    progress_bar = ttk.Progressbar(progress_frame, variable=progress_var,
                                   maximum=100, length=820)
    progress_bar.pack(fill=tk.X)

    convert_frame = tk.Frame(root, bg="#f8f9fa")
    convert_frame.pack(pady=10)
    convert_btn = ttk.Button(convert_frame, text="🚀 开始转换",
                             style="Big.TButton", width=30)
    convert_btn.pack()
    convert_btn.config(command=lambda: start_convert(var, root, progress_var,
                                                     convert_btn, name_entry, file_path))

    root.mainloop()